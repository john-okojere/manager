
from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.http import HttpResponse

from .models import Inventory
from .forms import InventoryForm, StartDayForm, EndDayForm
from users.models import CustomUser
from .models import Day
from reportlab.pdfgen import canvas
from openpyxl import Workbook
from docx import Document

from .forms import StartDayForm

@login_required
def start_day(request):
    # Check if there is any day that has not ended
    unfinished_day = Day.objects.filter(staff=request.user, end=False).first()
    if unfinished_day:
            return redirect('arcade_cashier')

    if request.method == 'POST':
        form = StartDayForm(request.POST)
        if form.is_valid():
            day = form.save(commit=False)
            day.staff = request.user
            day.start_time = datetime.now().time()
            day.save()
            return redirect('arcade_cashier')
    else:
        form = StartDayForm()

    return render(request, 'arcade/start_day.html', {'form': form})


@login_required
def cashier(request):
    # Check if there is any day that has not ended
    unfinished_day = Day.objects.filter(staff=request.user, end=False).first()
    if not unfinished_day:
        return redirect('start_day')  # Redirect to the start day view if no day has started

    items = Inventory.objects.all()
    context = {
        'items': items
    }
    return render(request, 'arcade/pos.html', context)

@login_required
def add_inventory(request):
    # Check if the user has an associated staff profile
    try:
        user = request.user
    except AttributeError:
        messages.error(request, "You don't have permission to upload inventory.")
        return redirect('inventory_list')

    if request.method == "POST":
        form = InventoryForm(request.POST)
        if form.is_valid():
            # Save the form but don't commit to the database yet
            inventory_item = form.save(commit=False)
            # Associate the logged-in staff with the inventory
            inventory_item.staff = user
            inventory_item.save()
            messages.success(request, "Inventory item added successfully!")
            return redirect('inventory_list')
    else:
        form = InventoryForm()

    return render(request, 'arcade/inventory/add.html', {'form': form})

def inventory_list(request):
    items = Inventory.objects.all().order_by('-date')
    return render(request, 'arcade/inventory/list.html', {'items': items})


# Update Inventory View
def update_inventory(request, pk):
    item = get_object_or_404(Inventory, pk=pk)
    if request.method == "POST":
        form = InventoryForm(request.POST, instance=item)
        if form.is_valid():
            form.save()
            # Respond to AJAX request with success message
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({"success": True, "message": "Inventory item updated successfully!"})
            messages.success(request, "Inventory item updated successfully!")
            return redirect('inventory_list')
        else:
            if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
                return JsonResponse({"success": False, "errors": form.errors})
    else:
        form = InventoryForm(instance=item)
    return render(request, 'arcade/inventory/update_inventory.html', {'form': form})


# Delete Inventory View
def delete_inventory(request, pk):
    item = get_object_or_404(Inventory, pk=pk)
    if request.method == "POST":
        item.delete()
        # Respond to AJAX request with success message
        if request.headers.get('X-Requested-With') == 'XMLHttpRequest':
            return JsonResponse({"success": True, "message": "Inventory item deleted successfully!"})
        messages.success(request, "Inventory item deleted successfully!")
        return redirect('inventory_list')
    return render(request, 'arcade/inventory/delete_inventory.html', {'item': item})


from django.http import JsonResponse
from .models import Sale, SaleItem, Inventory

def create_sale(request):
    unfinished_day = Day.objects.filter(staff=request.user, end=False).first()
    if request.method == 'POST':
        cashier_id = request.user.id
        total_amount  = 0
        cashier = CustomUser.objects.get(id=cashier_id)
        sale = Sale.objects.create(cashier=cashier, total=total_amount, day=unfinished_day)
        
        print(sale.total)
        return JsonResponse({'status': 'success', 'sale_id': sale.id})

from django.db.models import F

def add_sale_item(request):
    if request.method == 'POST':
        sale_id = request.POST.get('sale_id')
        product_id = request.POST.get('product_id')
        quantity = int(request.POST.get('quantity'))
        price = float(request.POST.get('price'))
        total = quantity * price
        print(quantity)
        # Fetch the sale and product
        sale = Sale.objects.get(id=sale_id)
        product = Inventory.objects.get(id=product_id)

        # Create the sale item
        sale_item = SaleItem.objects.create(
            sale=sale,
            product=product,
            quantity=quantity,
            price=price,
            total=total
        )

        # Update the sale total
        updated_total = SaleItem.objects.filter(sale=sale).aggregate(total_sum=Sum('total'))['total_sum'] or 0
        sale.total  = updated_total
        sale.save()

        # Refresh the sale object to get the updated total
        sale.refresh_from_db()

        return JsonResponse({
            'status': 'success',
            'sale_item_id': sale_item.id,
            'updated_sale_total': float(sale.total)
        })
def complete_sale(request):
    if request.method == 'POST':
        sale_id = request.POST.get('sale_id')
        print(sale_id)
        
        try:
            sale = Sale.objects.get(id=sale_id)
            sale.completed = True
            sale.save()
            
            return JsonResponse({'status': 'success', 'message': 'Sale marked as completed.'})
        except Sale.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Sale not found.'})
    return JsonResponse({'status': 'error', 'message': 'Invalid request method.'})

from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from .models import Sale, SaleItem, SaleDiscount, SaleItemDiscount, CustomUser

def apply_sale_discount(request):
    if request.method == 'POST':
        cashier = request.user
        sale_id = request.POST.get('sale_id')
        proposed_discount = request.POST.get('proposed_discount')

        sale = get_object_or_404(Sale, id=sale_id)

        # Create SaleDiscount
        SaleDiscount.objects.create(
            cashier=cashier,
            sale=sale,
            proposed_discount=proposed_discount,
        )

        return JsonResponse({'success': True, 'message': 'Sale discount created successfully!'})

    return JsonResponse({'success': False, 'message': 'Invalid request method.'})

def apply_sale_item_discount(request):
    if request.method == 'POST':
        cashier = request.user
        sale_item_id = request.POST.get('sale_item_id')
        proposed_discount = request.POST.get('proposed_discount')

        sale_item = get_object_or_404(SaleItem, id=sale_item_id)

        # Create SaleItemDiscount
        SaleItemDiscount.objects.create(
            cashier=cashier,
            sale=sale_item,
            proposed_discount=proposed_discount,
        )

        return JsonResponse({'success': True, 'message': 'Sale item discount created successfully!'})

    return JsonResponse({'success': False, 'message': 'Invalid request method.'})

from django.db.models import Sum
from django.db.models.functions import TruncDay, TruncWeek
from django.core.paginator import Paginator
from django.core.cache import cache
from django.db.models import Sum
from django.db.models.functions import TruncDay

def sales_history(request):
    if request.user.role == "Cashier":
        day = Day.objects.filter(staff=request.user, end=False).first()
        if not day:
            return redirect('start_day')

        sales = Sale.objects.filter(paid=True, cashier=request.user, day=day).only('id', 'cashier__username', 'completed', 'total', 'paid', 'date')
    elif request.user.role == "Manager" or request.user.level >= 3:
        sales = Sale.objects.filter(paid=True).only('id', 'cashier__username', 'completed', 'total', 'paid', 'date')
    else:
        sales = Sale.objects.none()

    paginator = Paginator(sales, 50)  # Show 50 sales per page
    page_number = request.GET.get('page')
    page_sales = paginator.get_page(page_number)

    sales_data = [
        {
            'id': sale.id,
            'cashier': sale.cashier.username,
            'completed': sale.completed,
            'total': float(sale.total),
            'paid': sale.paid,
            'date': sale.date.isoformat(),
        }
        for sale in page_sales
    ]

    cache_key = f'graph_data_{request.user.id}'
    graph_data = cache.get(cache_key)

    if not graph_data:
        graph_data = (
            sales
            .annotate(dates=TruncDay('date'))
            .values('dates')
            .annotate(total=Sum('total'))
            .order_by('dates')
        )
        cache.set(cache_key, list(graph_data), 3600)  # Cache for 1 hour

    graph_labels = [entry['dates'].strftime('%Y-%m-%d') for entry in graph_data]
    graph_values = [float(entry['total']) for entry in graph_data]

    if request.headers.get('x-requested-with') == 'XMLHttpRequest':
        return JsonResponse({'sales': sales_data, 'graph_labels': graph_labels, 'graph_values': graph_values})

    return render(request, 'arcade/sales_history.html', {
        'sales': sales_data,
        'graph_labels': graph_labels,
        'graph_values': graph_values,
        'page_sales': page_sales,
    })

from django.db.models import Sum
from django.db.models.functions import TruncDay
from django.db.models import Sum



from django.http import JsonResponse
from django.db.models import Sum
from django.db.models.functions import TruncDay
from .models import Sale

def fetch_sales_data(request):
    if request.method == 'GET':
        # Aggregate sales data
        graph_data = (
            Sale.objects.filter(completed=True)
            .annotate(day=TruncDay('date'))
            .values('day')
            .annotate(total=Sum('total'))
            .order_by('day')
        )

        graph_labels = [entry['day'].strftime('%Y-%m-%d') for entry in graph_data]
        graph_values = [float(entry['total']) for entry in graph_data]

        return JsonResponse({
            'graph_labels': graph_labels,
            'graph_values': graph_values,
        })

    return JsonResponse({'error': 'Invalid request method.'}, status=400)


from django.shortcuts import render
from django.http import JsonResponse
from .models import Sale  # Replace with your actual model for transactions
from django.db.models import Sum, Count
import datetime



# Generate a custom report (optional extension for downloadable formats like PDF/Excel)
def custom_report(request):
    # Handle filters and dynamic reports (e.g., by cashier, status, date, etc.)
    # Return processed data for export
    return JsonResponse({'status': 'Custom report generation in progress...'})

from django.db.models import Sum, Count
from decimal import Decimal
from django.shortcuts import render
from .models import Day, Payment, Sale, SaleItem, SaleDiscount, SaleItemDiscount, Inventory


def end_of_day_report(request):
    # Get the active day
    day = Day.objects.filter(staff=request.user, end=False).first()

    if not day:
        return redirect('start_day')  # Redirect to the start day view if no day has started

    # Calculate metrics
    sales = Sale.objects.filter(day=day)
    total_sales = sales.aggregate(total=Sum('total'))['total'] or Decimal('0.00')
    total_completed_sales = sales.filter(completed=True).count()
    total_pending_sales = sales.filter(completed=False).count()

    payments = Payment.objects.filter(sale__day=day)
    total_cash_payments = payments.filter(payment_type='cash').aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    total_card_payments = payments.filter(payment_type='card').aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    total_transfer_payments = payments.filter(payment_type='transfer').aggregate(total=Sum('amount'))['total'] or Decimal('0.00')

    total_cash_received = total_cash_payments
    total_payments = payments.aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    total_change_given = total_payments - total_sales
    expected_cash_at_hand = day.start_amount + total_cash_received - total_change_given

    # Discounts
    total_sale_discounts = SaleDiscount.objects.filter(sale__day=day).aggregate(total=Sum('proposed_discount'))['total'] or Decimal('0.00')
    total_approved_discounts = SaleDiscount.objects.filter(sale__day=day, approved=True).aggregate(total=Sum('proposed_discount'))['total'] or Decimal('0.00')

    # Inventory impact
    inventory_impact = (
        SaleItem.objects.filter(sale__day=day)
        .values('product__name')
        .annotate(total_quantity=Sum('quantity'), total_revenue=Sum('total'))
    )

    context = {
        'day': day,
        'total_sales': total_sales,
        'total_completed_sales': total_completed_sales,
        'total_pending_sales': total_pending_sales,
        'total_cash_payments': total_cash_payments,
        'total_card_payments': total_card_payments,
        'total_transfer_payments': total_transfer_payments,
        'total_cash_received': total_cash_received,
        'total_change_given': total_change_given,
        'expected_cash_at_hand': expected_cash_at_hand,

        'total_sale_discounts': total_sale_discounts,
        'total_approved_discounts': total_approved_discounts,
        'inventory_impact': inventory_impact,
    }

    return render(request, 'arcade/end-of-day.html', context)



import uuid
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.shortcuts import render
from .models import Payment

@csrf_exempt
def make_payment(request):
    if request.method == "POST":
        data = request.POST
        sale_id = data.get("sale_id")
        amount = data.get("amount")
        payment_method = data.get("payment_method")
        paid_by = data.get("paid_by")

        # For card payments, expect additional data
        if payment_method == "card":
            payment_id = uuid.uuid4().hex
            # Call Paystack API here to verify payment
            # Add your Paystack verification logic (e.g., check transaction ID)
        else:
            payment_id = None
        sale = Sale.objects.get(id = sale_id)
        # Save payment record
        payment = Payment.objects.create(
            sale = sale,
            cashier = request.user,
            amount=amount,
            payment_type=payment_method,
            paid_by=paid_by,
            payment_id=payment_id,
        )
        payment.save()

        return JsonResponse({"status": "success", "payment_id": payment.id})
    return JsonResponse({"status": "error", "message": "Invalid request"}, status=400)


def pay_for_sale(request):
    if request.method == 'POST':
        sale_id = request.POST.get('sale_id')
        print(sale_id)
        
        try:
            sale = Sale.objects.get(id=sale_id)
            sale.paid = True
            sale.save()
            
            return JsonResponse({'status': 'success', 'message': 'Sale marked as completed.'})
        except Sale.DoesNotExist:
            return JsonResponse({'status': 'error', 'message': 'Sale not found.'})

from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
import json

@csrf_exempt
def sale_receipt(request, sale_id):
    sale = get_object_or_404(Sale, id= sale_id)
    sale_items = sale.items.all()
    print("dcsxz")
    arcade_payment = Payment.objects.filter(sale=sale).last()
    print(arcade_payment)

    receipt_data = {
        "id": sale.id,
        "cashier": sale.cashier.username,
        "cashier_id": sale.cashier.id,
        "date": sale.date.strftime("%Y-%m-%d %H:%M:%S"),
        "total": sale.total,
        "items": [
            {
                "product": item.product.name,
                "quantity": item.quantity,
                "price": item.price,
                "total": item.total
            }
            for item in sale_items
        ],
        'payment':{
            'id': arcade_payment.id,
            'amount': arcade_payment.amount,
            'change': arcade_payment.amount - arcade_payment.sale.total,
            'by': arcade_payment.paid_by,
            'type': arcade_payment.payment_type,
        }
    }
    return JsonResponse(receipt_data)


from django.http import JsonResponse
from django.shortcuts import get_object_or_404
from django.views.decorators.http import require_GET, require_POST
from .models import Sale
from datetime import datetime

# API to fetch sales data
@require_GET
def fetch_eod_sales_data(request):
    if not request.user.is_authenticated:
        return JsonResponse({"error": "Unauthorized"}, status=401)

    # Check if the cashier has ended the day
    if not request.user.CustomUser.has_ended_day:
        return JsonResponse({"error": "End of day process not completed."}, status=403)

    # Filter sales for today
    today = datetime.now().date()
    sales = Sale.objects.filter(date__date=today)
    
    data = [
        {
            "id": sale.id,
            "cashier": str(sale.cashier),
            "total": float(sale.total),
            "completed": "Yes" if sale.completed else "No",
            "paid": float(sale.paid),
            "date": sale.date.strftime("%Y-%m-%d %H:%M:%S"),
            "payment_method": sale.payment_method if hasattr(sale, 'payment_method') else "N/A",
        }
        for sale in sales
    ]
    return JsonResponse(data, safe=False)

# API to end the day
def end_day(request):
    if not request.user.is_authenticated:
        return JsonResponse({"error": "Unauthorized"}, status=401)
    
    unfinished_day = Day.objects.filter(staff=request.user, end=False).first()
    if not unfinished_day:
        return redirect('start_day')  # Redirect to the start day view if no day has started
    
    unfinished_day.end_time  = datetime.now().time()
    unfinished_day.end = True
    unfinished_day.no_of_sales = Sale.objects.filter(day = unfinished_day).count()
    unfinished_day.end_amount = Sale.objects.filter(day = unfinished_day).aggregate(Sum('total'))['total__sum']
    unfinished_day.save()
    return redirect('end_of_day_report')  # Redirect to the start day view if no day has started

# Reset 'end day' status at midnight (Optional: handled via management command or cron job)

def audit_report(request):
    # Aggregate data
    total_sales = Sale.objects.filter(completed=True).aggregate(total=Sum('total'))['total'] or 0
    total_items_sold = SaleItem.objects.aggregate(total=Sum('quantity'))['total'] or 0
    inventory_data = Inventory.objects.all()
    sales_per_day = (
        Sale.objects.filter(completed=True)
        .values('date__date')
        .annotate(total_sales=Sum('total'))
        .order_by('-date__date')
    )

    # Data for charts
    # inventory_chart_data = {
    #     "labels": [item.name for item in inventory_data],
    #     "values": [item.quantity for item in inventory_data],
    # }

    sales_chart_data = {
        "labels": [str(day['date__date']) for day in sales_per_day],
        "values": [float(day['total_sales']) for day in sales_per_day],
    }

    context = {
        "total_sales": total_sales,
        "total_items_sold": total_items_sold,
        "inventory_data": inventory_data,
        "sales_per_day": sales_per_day,
        "sales_chart_data": sales_chart_data,
    }
    return render(request, "arcade/audit_report.html", context)

from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required, user_passes_test
from django.contrib import messages
from .models import Day, SaleDiscount, SaleItemDiscount


@login_required
def approve_day(request, day_id):
    day = get_object_or_404(Day, id=day_id, end=True)
    if request.method == "POST":
        day.approved = True
        day.save()
        print(day)
        messages.success(request, f"Day {day.id} approved successfully.")
        return redirect("manage_approvals")
    # Calculate metrics
    sales = Sale.objects.filter(day=day)
    total_sales = sales.aggregate(total=Sum('total'))['total'] or Decimal('0.00')
    total_completed_sales = sales.filter(completed=True).count()
    total_pending_sales = sales.filter(completed=False).count()

    payments = Payment.objects.filter(sale__day=day)
    total_cash_payments = payments.filter(payment_type='cash').aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    total_card_payments = payments.filter(payment_type='card').aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    total_transfer_payments = payments.filter(payment_type='transfer').aggregate(total=Sum('amount'))['total'] or Decimal('0.00')

    total_cash_received = total_cash_payments
    total_payments = payments.aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    total_change_given = total_payments - total_sales
    expected_cash_at_hand = day.start_amount + total_cash_received - total_change_given

    # Discounts
    total_sale_discounts = SaleDiscount.objects.filter(sale__day=day).aggregate(total=Sum('proposed_discount'))['total'] or Decimal('0.00')
    total_approved_discounts = SaleDiscount.objects.filter(sale__day=day, approved=True).aggregate(total=Sum('proposed_discount'))['total'] or Decimal('0.00')

    # Inventory impact
    inventory_impact = (
        SaleItem.objects.filter(sale__day=day)
        .values('product__name')
        .annotate(total_quantity=Sum('quantity'), total_revenue=Sum('total'))
    )

    context = {
        'day': day,
        'total_sales': total_sales,
        'total_completed_sales': total_completed_sales,
        'total_pending_sales': total_pending_sales,
        'total_cash_payments': total_cash_payments,
        'total_card_payments': total_card_payments,
        'total_transfer_payments': total_transfer_payments,
        'total_cash_received': total_cash_received,
        'total_change_given': total_change_given,
        'expected_cash_at_hand': expected_cash_at_hand,

        'total_sale_discounts': total_sale_discounts,
        'total_approved_discounts': total_approved_discounts,
        'inventory_impact': inventory_impact,
        "day": day
    }
    return render(request, "arcade/approvals/approve_day.html", context)


@login_required
def approve_discount(request, discount_id):
    discount = get_object_or_404(SaleDiscount, id=discount_id, approved=False)
    discount_percentage = (discount.proposed_discount / discount.sale.total) * 100 if discount.sale.total else 0
    if request.method == "POST":
        discount.approved = True
        discount.approved_by = request.user
        discount.save()
        messages.success(request, f"Discount for sale {discount.sale.id} approved successfully.")
        return redirect("manage_approvals")
    return render(request, "arcade/approvals/approve_discount.html", {"discount": discount, 'discount_percentage': round(discount_percentage, 2)})


@login_required
def manage_approvals(request):
    pending_days = Day.objects.filter(approved=False, end=True)
    pending_discounts = SaleDiscount.objects.filter(approved=False)
    context = {
        "pending_days": pending_days,
        "pending_discounts": pending_discounts,
    }
    return render(request, "arcade/approvals/manage_approvals.html", context)

from django.shortcuts import render
from django.contrib.auth.decorators import login_required, user_passes_test
from .models import Day, Sale, SaleDiscount, Payment
from django.db.models.functions import TruncMonth
from django.utils.timezone import now

@login_required
def dashboard(request):
    # Get the current month and year
    current_month = now().month
    current_year = now().year

    # Filter data for the current month
    total_sales = Sale.objects.filter(
        completed=True,
        date__month=current_month,
        date__year=current_year
    ).count()

    total_discounts = SaleDiscount.objects.filter(
        approved=True,
        sale__date__month=current_month,
        sale__date__year=current_year
    ).count()

    pending_days = Day.objects.filter(
        end=True,
        end_time=None,
        date__month=current_month,
        date__year=current_year
    ).count()

    pending_discounts = SaleDiscount.objects.filter(
        approved=False,
        sale__date__month=current_month,
        sale__date__year=current_year
    ).count()

    total_payments = Sale.objects.filter(
        completed=True,
        date__month=current_month,
        date__year=current_year
    ).aggregate(total=Sum('total'))['total'] or 0

    # Recent data for the current month
    recent_days = Day.objects.filter(
        date__month=current_month,
        date__year=current_year
    ).order_by('-date')[:10]

    recent_sales = Sale.objects.filter(
        date__month=current_month,
        date__year=current_year
    ).order_by('-date')[:10]

    recent_discounts = SaleDiscount.objects.filter(
        sale__date__month=current_month,
        sale__date__year=current_year
    ).order_by('-sale__date')[:10]

    # Sales grouped by month for a chart
    sales_by_month = Sale.objects.filter(completed=True).annotate(
        month=TruncMonth('date')
    ).values('month').annotate(total=Sum('total')).order_by('month')

    discount_status = {
        'approved': SaleDiscount.objects.filter(
            approved=True,
            sale__date__month=current_month,
            sale__date__year=current_year
        ).count(),
        'pending': SaleDiscount.objects.filter(
            approved=False,
            sale__date__month=current_month,
            sale__date__year=current_year
        ).count(),
    }

    context = {
       
        "total_sales": total_sales,
        "total_discounts": total_discounts,
        "pending_days": pending_days,
        "pending_discounts": pending_discounts,
        "total_payments": total_payments,
        "recent_days": recent_days,
        "recent_sales": recent_sales,
        "recent_discounts": recent_discounts,
        "sales_by_month": sales_by_month,
        "discount_status": discount_status,
    }
    return render(request, "arcade/dashboard.html", context)

def daily_sales_report(request, start_date, end_date):
    sales = Sale.objects.filter(date__range=(start_date, end_date))
    total_sales = sales.aggregate(Sum('total'))['total__sum'] or 0
    num_sales = sales.count()
    avg_sale_value = total_sales / num_sales if num_sales > 0 else 0

    top_products = (
        SaleItem.objects.filter(sale__in=sales)
        .values('product__name')
        .annotate(total_sold=Sum('quantity'), total_revenue=Sum('total'))
        .order_by('-total_revenue')[:5]
    )

    staff_performance = (
        sales.values('cashier__username')
        .annotate(total_sales=Sum('total'), sales_count=Count('id'))
        .order_by('-total_sales')
    )

    return JsonResponse({
        'total_sales': total_sales,
        'num_sales': num_sales,
        'avg_sale_value': avg_sale_value,
        'top_products': list(top_products),
        'staff_performance': list(staff_performance),
    })


def inventory_sales_rate(request):
    inventory_data = (
        SaleItem.objects.values('product__name')
        .annotate(total_sold=Sum('quantity'), total_revenue=Sum('total'))
        .order_by('-total_sold')
    )

    return JsonResponse({
        'inventory_data': list(inventory_data),
    })


from django.http import JsonResponse
from datetime import datetime, timedelta
from .models import Sale, SaleItem

def sales_growth_graph_data(request):
    # Monthly Data
    monthly_data = (
        Sale.objects.filter(date__year=datetime.now().year)
        .annotate(month=TruncMonth('date'))
        .values('month')
        .annotate(total_sales=Sum('total'))
        .order_by('month')
    )

    # Weekly Data (Last 12 Weeks)
    today = datetime.now()
    weeks_ago = today - timedelta(weeks=12)
    weekly_data = (
        Sale.objects.filter(date__range=(weeks_ago, today))
        .annotate(week=TruncWeek('date'))
        .values('week')
        .annotate(total_sales=Sum('total'))
        .order_by('week')
    )

    # Daily Data (Last 30 Days)
    days_ago = today - timedelta(days=30)
    daily_data = (
        Sale.objects.filter(date__range=(days_ago, today))
        .annotate(day=TruncDay('date'))
        .values('day')
        .annotate(total_sales=Sum('total'))
        .order_by('day')
    )

    return JsonResponse({
        'monthly_data': list(monthly_data),
        'weekly_data': list(weekly_data),
        'daily_data': list(daily_data),
    })


def sales_by_staff_report(request):
    staff_performance = (
        Sale.objects.values('cashier__username')
        .annotate(total_sales=Sum('total'), sales_count=Count('id'))
        .order_by('-total_sales')
    )
    return JsonResponse({'staff_performance': list(staff_performance)})


def sales_by_product_report(request):
    product_performance = (
        SaleItem.objects.values('product__name')
        .annotate(total_units_sold=Sum('quantity'), total_revenue=Sum('total'))
        .order_by('-total_units_sold')
    )
    return JsonResponse({'product_performance': list(product_performance)})

def low_stock_report(request):
    low_stock_items = Inventory.objects.filter(stock_level__lte=10).values('name', 'stock_level')
    return JsonResponse({'low_stock_items': list(low_stock_items)})


def daily_cash_flow_report(request, start_date, end_date):
    payments = Payment.objects.filter(date__range=(start_date, end_date))
    payment_summary = payments.values('payment_type').annotate(total_amount=Sum('amount'))
    return JsonResponse({'cash_flow': list(payment_summary)})


def discounts_report(request):
    discounts = (
        SaleDiscount.objects.filter(approved=True)
        .values('cashier__username', 'sale__id')
        .annotate(total_discount=Sum('proposed_discount'))
        .order_by('-total_discount')
    )
    return JsonResponse({'discounts': list(discounts)})

from .models import Refund
def refunds_report(request):
    refunds = Refund.objects.values('sale__id', 'amount', 'reason', 'date').order_by('-date')
    return JsonResponse({'refunds': list(refunds)})


from django.shortcuts import render
from django.db.models import Sum, F, Count
from decimal import Decimal
from .models import Sale, SaleItem, Day, Payment

# Weekly Sales Report
def weekly_sales_report(request):
    sales = Sale.objects.filter(date__week=timezone.now().isocalendar()[1])
    total_sales = sales.aggregate(total=Sum('total'))['total'] or Decimal('0.00')
    return render(request, 'reports/weekly_sales_report.html', {'sales': sales, 'total_sales': total_sales})

# Monthly Sales Report
def monthly_sales_report(request):
    sales = Sale.objects.filter(date__month=timezone.now().month)
    total_sales = sales.aggregate(total=Sum('total'))['total'] or Decimal('0.00')
    return render(request, 'reports/monthly_sales_report.html', {'sales': sales, 'total_sales': total_sales})

# Inventory Turnover Report
def inventory_turnover_report(request):
    inventory_data = (
        SaleItem.objects.values('product__name')
        .annotate(total_sold=Sum('quantity'), revenue=Sum('total'))
        .order_by('-revenue')
    )
    return render(request, 'reports/inventory_turnover_report.html', {'inventory_data': inventory_data})

# Daily Cash Flow Report
def daily_cashflow_report(request):
    payments = Payment.objects.filter(date__date=timezone.now().date())
    payment_summary = (
        payments.values('payment_type')
        .annotate(total=Sum('amount'))
        .order_by('payment_type')
    )
    total_cashflow = payments.aggregate(total=Sum('amount'))['total'] or Decimal('0.00')
    return render(request, 'reports/daily_cashflow_report.html', {'payment_summary': payment_summary, 'total_cashflow': total_cashflow})

# Sales Tax Report
def sales_tax_report(request):
    sales = Sale.objects.filter(date__month=timezone.now().month)
    total_sales = sales.aggregate(total=Sum('total'))['total'] or Decimal('0.00')
    tax_rate = Decimal('0.075')  # Assuming 7.5% sales tax
    total_tax = total_sales * tax_rate
    return render(request, 'reports/sales_tax_report.html', {'total_sales': total_sales, 'total_tax': total_tax})

from django.db.models import Sum, Avg, Count
from datetime import timedelta
from django.utils.timezone import now

def api_sales_growth(request):
    # Calculate Monthly Sales
    monthly_sales = (
        Sale.objects.filter(date__year=now().year)
        .annotate(month=TruncMonth('date'))
        .values('month')
        .annotate(total_sales=Sum('total'))
        .order_by('month')
    )

    # Calculate Weekly Sales
    weekly_sales = (
        Sale.objects.filter(date__gte=now() - timedelta(weeks=4))
        .annotate(week=TruncMonth('date'))
        .values('week')
        .annotate(total_sales=Sum('total'))
        .order_by('week')
    )

    # Calculate Daily Sales
    daily_sales = (
        Sale.objects.filter(date__gte=now() - timedelta(days=30))
        .extra(select={'day': "DATE(date)"})
        .values('day')
        .annotate(total_sales=Sum('total'))
        .order_by('day')
    )

    # Inventory Sales Rate
    inventory_sales = (
        SaleItem.objects.values('product__name')
        .annotate(total_sold=Sum('quantity'))
        .order_by('-total_sold')
    )

    data = {
        "monthly_data": list(monthly_sales),
        "weekly_data": list(weekly_sales),
        "daily_data": list(daily_sales),
        "inventory_data": list(inventory_sales),
    }
    return JsonResponse(data)

def api_sales_performance(request):
    # Sales Performance by Staff
    sales_performance = (
        Sale.objects.values('cashier__username')
        .annotate(
            total_sales=Sum('total'),
            sales_count=Count('id'),
            avg_sale_value=Avg('total')
        )
        .order_by('-total_sales')
    )

    # Top Products
    top_products = (
        SaleItem.objects.values('product__name')
        .annotate(
            total_revenue=Sum('total'),
            units_sold=Sum('quantity')
        )
        .order_by('-total_revenue')[:10]
    )

    # Payment Summary
    payment_summary = (
        Payment.objects.values('payment_type')
        .annotate(total_amount=Sum('amount'))
        .order_by('-total_amount')
    )

    data = {
        "sales_performance": list(sales_performance),
        "top_products": list(top_products),
        "payment_summary": list(payment_summary),
    }
    return JsonResponse(data)
# API: Sales by Product
def api_sales_by_product(request):
    product_sales = (
        SaleItem.objects.values('product__name')
        .annotate(total_revenue=Sum('total'), total_quantity=Sum('quantity'))
        .order_by('-total_revenue')
    )
    return JsonResponse({'sales_by_product': list(product_sales)})

# API: Payment Summary
def api_payment_summary(request):
    payment_summary = (
        Payment.objects.values('payment_type')
        .annotate(total_amount=Sum('amount'))
        .order_by('payment_type')
    )
    return JsonResponse({'payment_summary': list(payment_summary)})

def export_pdf(request):
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="sales_report.pdf"'
    p = canvas.Canvas(response)
    p.drawString(100, 800, "Comprehensive Sales Report")
    
    # Fetch sales data
    sales = Sale.objects.all()
    y = 750
    for sale in sales:
        p.drawString(100, y, f"Sale ID: {sale.id}, Total: {sale.total}, Date: {sale.date.strftime('%Y-%m-%d %H:%M:%S')}")
        y -= 20
        if y < 50:
            p.showPage()
            y = 800

    p.showPage()
    p.save()
    return response
# Export Excel
def export_excel(request):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet')
    response['Content-Disposition'] = 'attachment; filename="sales_report.xlsx"'
    wb = Workbook()
    ws = wb.active
    ws.title = "Sales Report"
    # Add your content dynamically
    ws.append(["Header 1", "Header 2", "Header 3"])  # Example
    ws.append(["Data 1", "Data 2", "Data 3"])  # Example
    wb.save(response)
    return response

# Export Doc
def export_doc(request):
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = 'attachment; filename="sales_report.docx"'
    doc = Document()
    doc.add_heading('Comprehensive Sales Report', 0)
    # Add your content dynamically
    doc.add_paragraph("This is a sample report generated dynamically.")
    doc.save(response)
    return response


